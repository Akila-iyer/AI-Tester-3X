"""Tool: Extract text from PDF, DOCX, XLSX, CSV files."""

import csv
import os
from typing import List, Tuple

import pdfplumber
import docx
import openpyxl


def extract(filepath: str) -> Tuple[List[str], str]:
    """
    Extract content from a file as a list of text lines.
    Returns (lines, file_type_identifier).
    """
    ext = os.path.splitext(filepath)[1].lower()

    if ext == ".pdf":
        return _extract_pdf(filepath), "pdf"
    elif ext == ".docx":
        return _extract_docx(filepath), "docx"
    elif ext == ".xlsx" or ext == ".xls":
        return _extract_xlsx(filepath), "xlsx"
    elif ext == ".csv":
        return _extract_csv(filepath), "csv"
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def _extract_pdf(filepath: str) -> List[str]:
    lines = []
    with pdfplumber.open(filepath) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                lines.extend(text.split("\n"))
    return lines


def _extract_docx(filepath: str) -> List[str]:
    lines = []
    doc = docx.Document(filepath)

    # Paragraphs
    for para in doc.paragraphs:
        lines.append(para.text)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text for cell in row.cells)
            lines.append(row_text)

    return lines


def _extract_xlsx(filepath: str) -> List[str]:
    lines = []
    wb = openpyxl.load_workbook(filepath, data_only=True)
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        lines.append(f"[Sheet: {sheet_name}]")
        for row in ws.iter_rows(values_only=True):
            row_text = " | ".join(
                str(cell) if cell is not None else "" for cell in row
            )
            lines.append(row_text)
    wb.close()
    return lines


def _extract_csv(filepath: str) -> List[str]:
    lines = []
    encoding = "utf-8"
    try:
        with open(filepath, "r", encoding=encoding) as f:
            reader = csv.reader(f)
            for row in reader:
                lines.append(" | ".join(row))
    except UnicodeDecodeError:
        # Fallback to latin-1
        with open(filepath, "r", encoding="latin-1") as f:
            reader = csv.reader(f)
            for row in reader:
                lines.append(" | ".join(row))
    return lines
